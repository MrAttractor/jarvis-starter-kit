# -*- coding: utf-8 -*-
"""Construit les courriers de demande de partenariat, prets a envoyer.

Regle de ce script, apprise a la relecture du 21/08 : il ne produit PAS de
modele. Il produit une lettre finie par entreprise, nom rempli, date posee,
reference attribuee. **Aucun crochet ne doit survivre dans un document qui part
chez un destinataire.** Un controle automatique le verifie avant l'ecriture.

Le texte de reference est celui valide par Mac Arthur sur le courrier
TAP TAP SEND du 21 aout 2026. On n'y touche plus : seuls changent le nom de
l'entreprise et le numero de reference. La variante champagne est la seule
exception, pour une raison de droit expliquee dans son bloc.

Usage :
    python construire-lettre.py              tous les courriers
    python construire-lettre.py orange-money un seul
"""
import sys
import os
import re
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                '..', '..', '..', '..', 'scripts'))
from docx_commun import nouveau_document, inline, GRIS  # noqa: E402

ICI = os.path.dirname(os.path.abspath(__file__))
SORTIE = os.path.join(ICI, 'a-envoyer')
DATE = 'Rennes, le 21 août 2026'

# --- Le corps commun, texte valide le 21/08 ---------------------------------

SALUT = "Madame, Monsieur,"

SOCLE = [
 "Le Festival des Grillades est né à Abidjan en 2008 d'une intuition simple : la gastronomie est un langage universel, capable de rassembler, de faire rayonner une culture et de créer de la valeur économique. Il s'est tenu chaque année depuis, sans discontinuité, et la dix-huitième édition abidjanaise s'est close en 2025.",

 "En 2026, le Festival franchit un cap : sept éditions, trois pays, de Cotonou à Boundiali, d'Abidjan à Dakar. Et, le 11 octobre prochain à Bobigny, sa deuxième édition parisienne.",

 "Paris est une vitrine stratégique. C'est là que se retrouvent la diaspora ivoirienne et un public international curieux d'expériences authentiques. La première édition parisienne a affiché complet. Cette année, l'événement se tiendra de midi à vingt et une heures, dans une salle de cinq cents places, et nous attendons quatre cents visiteurs qui seront présents pour cette célébration.",
]

POLITESSE = ("En espérant que vous donnerez une suite favorable à notre demande "
             "d'audience, veuillez agréer, Madame, Monsieur, l'expression de "
             "notre haute considération.")


def sponsoring(nom):
    return [
     "Advantage Conseils, propriétaire du Festival, nous a confié la production exécutive de cette édition. À ce titre, nous constituons le cercle de ses partenaires en France, et nous venons respectueusement solliciter %s pour s'y associer." % nom,

     "S'associer au Festival, c'est être présent là où votre public se réunit de lui-même, en famille, pour une journée entière : un événement positif et festif, où votre marque n'est pas affichée mais vécue.",

     "Nous joignons à ce courrier le dossier de présentation de l'édition parisienne, qui vous donnera plus d'informations sur l'événement, son public et les dispositifs de visibilité possibles. Les modalités d'un partenariat se construisent ensuite avec vous, selon ce que vous souhaitez y faire.",
    ]


def collaboration(nom):
    """Variante des maisons de champagne.

    L'article L.3323-2 du code de la sante publique interdit le parrainage d'un
    evenement par un producteur de boissons alcooliques. Ce courrier ne demande
    donc pas un partenariat : il presente l'opportunite d'une collaboration
    commerciale et ouvre la discussion. Aucune contrepartie de visibilite n'y
    est proposee, nulle part.
    """
    return [
     "Quatre cents convives attendus, une table à l'honneur du début à la fin, un bar ouvert de midi à vingt et une heures : la carte de cette édition se construit en ce moment, et nous souhaitons y associer une grande maison de champagne.",

     "Advantage Conseils, propriétaire du Festival, nous a confié la production exécutive de cette édition. C'est à ce titre que nous venons vers %s, pour vous présenter l'opportunité d'une collaboration autour de cette journée et, si le principe vous intéresse, ouvrir la discussion sur la forme qu'elle pourrait prendre." % nom,

     "Nous pensons à un accord d'approvisionnement, à l'achat ou en dépôt-vente, le service étant assuré par notre équipe et le relevé des volumes réellement servis vous étant remis après l'événement. Nous restons ouverts à toute autre forme qui vous conviendrait mieux.",

     "Précisons le cadre : notre démarche est strictement commerciale. Nous ne sollicitons aucun parrainage, que la réglementation française interdit pour les boissons alcooliques, et n'attendons aucune contrepartie publicitaire. Vous trouverez ci-joint le dossier de présentation de l'édition.",
    ]


OBJET_SPONSO = ("Sponsoring du Festival des Grillades de Paris, "
                "2e édition, 11 octobre 2026")
OBJET_COLLAB = ("Édition parisienne du Festival des Grillades, "
                "11 octobre 2026, proposition de collaboration")

LETTRES = [
 # TAP TAP SEND : courrier etabli et envoye par Mac Arthur le 21/08, reference
 # 08.26-01. Sa version fait foi, on ne la regenere pas pour ne pas creer un
 # second exemplaire du meme courrier.
 ('orange-money',   'ORANGE MONEY FRANCE', '08.26-02', OBJET_SPONSO, sponsoring),
 ('veuve-clicquot', 'VEUVE CLICQUOT',      '08.26-03', OBJET_COLLAB, collaboration),
 ('mercier',        'CHAMPAGNE MERCIER',   '08.26-04', OBJET_COLLAB, collaboration),
 ('moet-chandon',   'MOËT & CHANDON',      '08.26-05', OBJET_COLLAB, collaboration),
]

PIED = ("THIM PRODUCTION · SASU au capital de 500 € · 22 rue Charles Duclos, 35000 Rennes\n"
        "SIREN 953 657 251 R.C.S. Rennes · SIRET 953 657 251 000 13\n"
        "thimproduction22@gmail.com · 06 12 87 94 04")


def para(doc, texte='', taille=11, gras=False, align=None, apres=6,
         couleur=None, interligne=1.15):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(apres)
    p.paragraph_format.line_spacing = interligne
    if align is not None:
        p.alignment = align
    if texte:
        inline(p, texte, base_gras=gras, base_taille=taille)
        for r in p.runs:
            # Jamais de surlignage dans un courrier : Word le conserve quand on
            # tape par-dessus, et la lettre partirait avec des marques fluo.
            r.font.highlight_color = None
            if couleur:
                r.font.color.rgb = couleur
    return p


def construire(cle, nom, ref, objet, corps_fn):
    doc = nouveau_document('Courrier de partenariat, Festival des Grillades de Paris',
                           auteur='Thim Production')
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.4)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    logo = os.path.join(ICI, 'logos', 'thim.png')
    if os.path.exists(logo):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.add_run().add_picture(logo, height=Cm(1.35))

    para(doc, "Producteur délégué de l'édition parisienne du Festival des Grillades",
         taille=8.5, couleur=GRIS, apres=10)

    para(doc, DATE, taille=10.5, align=WD_ALIGN_PARAGRAPH.RIGHT, apres=12)

    for ligne in ("À l'attention de la Direction de", nom, 'France'):
        para(doc, ligne, taille=10.5, align=WD_ALIGN_PARAGRAPH.RIGHT, apres=1)

    para(doc, apres=10)
    para(doc, 'Réf. : FGP2026 / PART / ' + ref, taille=10, apres=2)
    para(doc, 'Objet : ' + objet, taille=10, gras=True, apres=12)

    for bloc in [SALUT] + SOCLE + corps_fn(nom) + [POLITESSE]:
        para(doc, bloc, taille=10, apres=6, interligne=1.08)

    para(doc, apres=2)
    para(doc, 'Arnaud YORO', gras=True, apres=0)
    para(doc, 'Manager général', taille=10, apres=0)
    para(doc, 'Thim Production', taille=10, apres=8)

    para(doc, "Pièce jointe : dossier de présentation de l'édition parisienne.",
         taille=9.5, couleur=GRIS, apres=0)

    pied = sec.footer.paragraphs[0]
    pied.text = ''
    pied.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, ligne in enumerate(PIED.split('\n')):
        if i:
            pied.add_run().add_break()
        r = pied.add_run(ligne)
        r.font.size = Pt(7.5)
        r.font.color.rgb = GRIS

    # Controle avant ecriture : aucun crochet ne doit rester.
    texte = '\n'.join(p.text for p in doc.paragraphs)
    restes = re.findall(r'\[[^\]]*\]', texte)
    if restes:
        raise SystemExit('CHAMPS NON REMPLIS dans %s : %s' % (cle, restes))

    os.makedirs(SORTIE, exist_ok=True)
    chemin = os.path.join(SORTIE, 'Courrier-%s.docx' % cle)
    try:
        doc.save(chemin)
    except PermissionError:
        print('VERROUILLE, non refait :', os.path.basename(chemin))
        return
    print('ecrit', os.path.basename(chemin))


def main():
    cles = sys.argv[1:]
    for l in LETTRES:
        if not cles or l[0] in cles:
            construire(*l)


if __name__ == '__main__':
    main()
