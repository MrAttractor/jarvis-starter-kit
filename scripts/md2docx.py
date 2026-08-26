# -*- coding: utf-8 -*-
"""Markdown -> Word modifiable.

Usage :
  python scripts/md2docx.py source.md sortie.docx [--titre "..."] [--sous "..."]
                            [--avert "..."] [--pied "..."] [--auteur "..."]
                            [--sauter-jusqua "## Titre de section"]

Le titre du document est pris sur la premiere ligne "# ..." si --titre est absent.
Les blocs entre triples accents graves deviennent des blocs a copier (lettre type,
message, script), dans la police du document et non en machine a ecrire.
"""
import re
import sys
import argparse

sys.path.insert(0, __file__.rsplit('\\', 1)[0].rsplit('/', 1)[0])
from docx.shared import Pt
from docx_commun import (nouveau_document, tableau, citation, filet,
                        bloc_a_copier, inline, pied_de_page)

SEP_TABLEAU = re.compile(r'^\s*\|[\s:|-]+\|\s*$')
SANS_NOTES = False


def cellules(ligne):
    return [c.strip() for c in ligne.strip().strip('|').split('|')]


def convertir(md, doc):
    lignes = md.split('\n')
    i = 0
    while i < len(lignes):
        ln = lignes[i]

        if not ln.strip():
            i += 1
            continue

        if ln.strip() == '---':
            filet(doc)
            i += 1
            continue

        # Bloc a copier
        if ln.strip().startswith('```'):
            i += 1
            buf = []
            while i < len(lignes) and not lignes[i].strip().startswith('```'):
                buf.append(lignes[i])
                i += 1
            i += 1
            while buf and not buf[0].strip():
                buf.pop(0)
            while buf and not buf[-1].strip():
                buf.pop()
            bloc_a_copier(doc, buf)
            continue

        # Tableau
        if ln.lstrip().startswith('|') and i + 1 < len(lignes) \
                and SEP_TABLEAU.match(lignes[i + 1]):
            entetes = cellules(ln)
            i += 2
            corps = []
            while i < len(lignes) and lignes[i].lstrip().startswith('|'):
                corps.append(cellules(lignes[i]))
                i += 1
            tableau(doc, entetes, corps)
            continue

        # Citation. Les notes de marge internes, reperees au chevron, ne sortent
        # jamais dans un document contractuel envoye a un tiers : elles disent
        # pourquoi une clause existe, ce qui est notre raisonnement, pas le sien.
        if ln.lstrip().startswith('>'):
            buf = []
            while i < len(lignes) and lignes[i].lstrip().startswith('>'):
                buf.append(lignes[i].lstrip()[1:].strip())
                i += 1
            texte = ' '.join(x for x in buf if x)
            if SANS_NOTES and ('▸' in texte or texte.startswith('**Version ')):
                continue
            citation(doc, texte)
            continue

        # Listes
        puce = re.match(r'^\s*[-*]\s+(.*)', ln)
        num = re.match(r'^\s*\d+\.\s+(.*)', ln)
        if puce or num:
            motif = r'^\s*\d+\.\s+(.*)' if num else r'^\s*[-*]\s+(.*)'
            style = 'List Number' if num else 'List Bullet'
            items, courant = [], None
            while i < len(lignes):
                m = re.match(motif, lignes[i])
                if m:
                    if courant is not None:
                        items.append(courant)
                    courant = m.group(1).strip()
                    i += 1
                elif lignes[i].strip() and lignes[i].startswith(('  ', '\t')) \
                        and courant is not None:
                    courant += ' ' + lignes[i].strip()
                    i += 1
                else:
                    break
            if courant is not None:
                items.append(courant)
            for it in items:
                p = doc.add_paragraph(style=style)
                p.paragraph_format.space_after = Pt(4)
                inline(p, it)
            continue

        # Titres
        m = re.match(r'^(#{2,4})\s+(.*)', ln)
        if m:
            niveau = min(len(m.group(1)) - 1, 3)
            p = doc.add_paragraph(style='Heading %d' % niveau)
            inline(p, m.group(2).strip(), base_gras=True)
            i += 1
            continue

        if ln.startswith('# '):
            i += 1
            continue

        # Paragraphe
        buf = []
        while i < len(lignes) and lignes[i].strip() \
                and not lignes[i].lstrip().startswith(('|', '>', '#', '```')) \
                and lignes[i].strip() != '---' \
                and not re.match(r'^\s*([-*]|\d+\.)\s+', lignes[i]):
            buf.append(lignes[i].strip())
            i += 1
        if buf:
            p = doc.add_paragraph()
            inline(p, ' '.join(buf))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('source')
    ap.add_argument('sortie')
    ap.add_argument('--titre')
    ap.add_argument('--sous')
    ap.add_argument('--avert')
    ap.add_argument('--pied')
    ap.add_argument('--auteur', default='Thim Production')
    ap.add_argument('--sauter-jusqua')
    a = ap.parse_args()

    md = open(a.source, encoding='utf-8').read()

    titre = a.titre
    if not titre:
        for ln in md.split('\n'):
            if ln.startswith('# '):
                titre = ln[2:].strip()
                break
    titre = titre or 'Document'

    if a.sauter_jusqua:
        pos = md.find(a.sauter_jusqua)
        if pos == -1:
            raise SystemExit('section introuvable : ' + a.sauter_jusqua)
        md = md[pos:]

    doc = nouveau_document(titre, a.sous, a.auteur, a.avert)
    convertir(md, doc)
    if a.pied:
        pied_de_page(doc, a.pied)
    doc.save(a.sortie)
    print('ecrit', a.sortie)


if __name__ == '__main__':
    main()
