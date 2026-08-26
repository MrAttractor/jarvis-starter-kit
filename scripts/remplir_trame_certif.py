# -*- coding: utf-8 -*-
"""Remplit la trame officielle de l'Academie Puissance Mentale.

Pourquoi ce script : la consigne impose "la trame fournie (format et police
d'origine)". On part donc du .docx de l'Academie et on ne fait qu'y verser du
texte, sans jamais recreer un document. La police n'est nulle part forcee : le
style Normal du fichier (Calibri) fait foi, on ne pose que la taille.

Usage : python scripts/remplir_trame_certif.py
"""
import copy
import os
import sys

from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL

DOSSIER = os.path.join('livrables', 'ecosysteme-attractor',
                       'certification-preparation-mentale')
SOURCE = os.path.join(DOSSIER, '2026-07 Support certification PM - Nouveau logo - Word.docx')
SORTIE = os.path.join(DOSSIER, 'RENDU-THEME-B-Eric-trinome.docx')

CANDIDATS = ['- Gisèle', '- Anouschka', '- Mac Arthur']

AXES = (
    "Sommet plat, 4 points seulement entre le 7 (33) et le 1 (29) : je pose une hypothèse de "
    "travail, que je valide avec lui en séance 2, jamais un verdict. Un profil fort n'est pas un "
    "profil mal utilisé, il a une zone de maîtrise et une zone critique : je ne fais pas baisser "
    "son 8, je le rééquilibre. **8 vers 2**, écouter le besoin de l'autre ; **3 vers 6**, fédérer "
    "et demander de l'aide ; **7 vers 5**, sobriété ; **1 vers 4**, écouter ses émotions. Les "
    "trois premiers sont aussi ses directions d'intégration : deux lectures différentes donnent "
    "le même plan de travail. Le 1, lui, ne s'intègre pas en 4, il s'y désintègre : son 4 faible "
    "est une vulnérabilité, pas un simple déséquilibre. Piège central : « je dois leur en donner "
    "beaucoup » est une croyance jamais testée, il n'a aucun fait sur ce que ses clients "
    "attendent, et **il se lance dans un métier de groupe avec le relationnel le plus bas de tout "
    "son profil (2 à 14)**. Il se prépare à une performance de contenu là où l'exercice est "
    "relationnel : sa peur est légitime mais porte sur le mauvais objet, et sa réponse, en donner "
    "plus, aggrave l'épuisement sans traiter la cause. Il touche **les deux zones critiques**, la "
    "haute en sur-activation (projet 24/24) et la basse en sous-activation (fatigue) : il lui "
    "faut de quoi descendre et de quoi remonter. Avec un 4 à 22, aucun voyant d'alarme interne, "
    "d'où le kiff avant le cash. Enfin le secret gardé à sa famille produit la même chose que sa "
    "fatigue et l'éloignement de ses proches : **il porte seul**. Vigilance : ce faisceau évoque "
    "aussi un épuisement, d'où un screening en séance 1 et une orientation médicale si des "
    "signaux cliniques apparaissent."
)

SEANCES = [
    ('1', '1h30',
     "**Thème : ce qui dépend de moi et ce qui n'en dépend pas.** Anamnèse en mode explorateur sur "
     "tous les domaines de vie, recueil de ses perceptions cerveau, cœur, comportement, tri des "
     "faits et des croyances, screening d'épuisement, puis objectif et 3 indicateurs posés aux "
     "échelles de 0 à 10. **EMCT : météo intérieure**, 7 plans, dont le relationnel, la "
     "responsabilité et la satisfaction de vie. Sortie sur « quelle est la plus petite des actions "
     "qui dépend que de vous ? »",
     "Noter chaque jour ses 3 indicateurs. Rien d'autre : la charge de départ est volontairement "
     "légère chez un homme en surcharge.",
     "Fixe l'objectif et sa mesure. Chez un 3 à 30 qui indexe sa valeur sur le résultat, "
     "distinguer ce qu'il contrôle de ce qu'il ne contrôle pas fait déjà baisser la pression."),

    ('2', '1h',
     "**Thème : un profil fort n'est pas un défaut ; zone de maîtrise et deux zones critiques.** "
     "Décryptage en échange : je valide avec lui ses forces, ses peurs, ses faiblesses, puis les 4 "
     "rééquilibrages, puis l'explication du déroulé du coaching et des orientations stratégiques. "
     "**EMCT : gestion des émotions par les respirations**, 3 cycles, avec comparaison à sa "
     "respiration du début de séance.",
     "Préparer sa roue : sur quoi veut-il travailler. Pratiquer la respiration une fois par jour, "
     "à heure fixe.",
     "Il comprend que sa fatigue vient d'un fonctionnement, pas d'un manque de compétence. Il "
     "repart avec le seul outil utilisable seul dès le lendemain, avant une session : indicateur "
     "1. Et un 8 a besoin de voir la carte avant de s'engager."),

    ('3', '1h',
     "**Thème : une croyance n'est pas un fait, et pourtant elle organise les comportements.** "
     "Lecture de sa roue, constat par lui d'abord, puis entonnoir qui part de tous ses domaines et "
     "resserre jusqu'à la croyance « je dois leur en donner beaucoup ». Recadrage : je l'accueille, "
     "je ne la contredis jamais, je cherche avec lui un moment où l'inverse s'est vérifié. **EMCT : "
     "ancrage de la croyance aidante trouvée.**",
     "Cinq solutions dont deux farfelues, il en choisit deux ou trois. Repositionne sa roue.",
     "La croyance traitée est celle qui organise toute sa surcharge. Et les actions sont choisies "
     "par lui, jamais imposées : un 8 à 32 n'exécute pas longtemps les consignes d'un autre."),

    ('4', '1h',
     "**Thème : les valeurs comme filtre de décision, et le kiff avant le cash.** Travail "
     "stratégique : 3 ou 4 personnes qu'il admire, les qualités qu'il leur reconnaît, 10 valeurs "
     "puis les 5 plus fortes, avec la question « après ma mort, qu'est-ce que je veux que mes "
     "proches disent de moi ? ». Puis sa définition de chacune en une phrase. **EMCT : précision "
     "des valeurs, puis les deux chemins** : un passage sur le chemin où elles ne sont pas "
     "nourries, deux sur l'autre, et on ferme dessus.",
     "Relire sa dernière session à la lumière de ses 5 valeurs et noter ce qui n'y correspond pas.",
     "Avec un 4 à 22 il n'a aucun voyant d'alarme interne : ce protocole est le voyant, il fait "
     "ressentir la trajectoire au lieu de l'expliquer. Ses valeurs deviennent le filtre qui tranche "
     "ce qu'il coupe, au lieu de tout donner."),

    ('5', '1h',
     "**Thème : on ne peut pas être à la hauteur d'une attente qu'on n'a jamais demandée.** Les "
     "quatre portes du besoin réel : problèmes, peurs, souhaits, désirs. Qui veut-il accompagner, "
     "ce qu'il croit qu'ils attendent, fait ou croyance, puis construction de sa trame d'entretien. "
     "**EMCT : position de l'autre** : il ressent ce que produit chez un participant le fait d'être "
     "vraiment entendu, puis futurisation sur son premier entretien.",
     "Dix entretiens d'écoute, réponses rapportées telles quelles. Consigne stricte : il ne vend "
     "rien, il écoute.",
     "Indicateur 2. Traite sa peur à la racine : tant qu'il devine, il ne peut que craindre. Et "
     "travaille son 2 à 14 sur le terrain même de son métier. Pour un 8 fort, se taire vingt "
     "minutes est l'exercice."),

    ('6', '1h',
     "**Thème : personne ne performe durablement seul, et il ne peut pas installer dans son groupe "
     "l'ouverture qu'il refuse pour lui-même.** Débrief des dix entretiens. Le secret familial, "
     "abordé ici parce que l'alliance est solide : que se passerait-il s'ils savaient, et surtout "
     "que te coûte-t-il de le porter seul ? Recadrage de « demander de l'aide, c'est avouer une "
     "faiblesse ». **EMCT : changement de zone**, descendre et remonter.",
     "En parler à une personne de confiance, une seule : exposition graduée. Solliciter une aide "
     "concrète sur un point précis de son lancement.",
     "Indicateur 3. C'est le même mouvement qui produit sa fatigue et l'éloignement de ses proches. "
     "Sans cette séance, un homme qui porte seul ne tient pas quatre mois."),

    ('7', '1h',
     "**Thème : moins mais plus profond, et choisir son couloir réduit le périmètre sur lequel il "
     "faut être à la hauteur.** Les 3 besoins qui reviennent, offre reconstruite autour d'eux, "
     "réduction à une page et un seul message, confrontation à ses 5 valeurs, positionnement dit en "
     "une phrase (pour qui il est le meilleur choix et pour qui il ne l'est pas), puis la question "
     "du prix. **EMCT : E.C.O.**, en balayant les domaines de vie car il débute : on cherche un "
     "état, pas un savoir-faire. Puis renforcement positif.",
     "Réécrire l'offre en une page et la tester auprès de 3 personnes interrogées. Construire sa "
     "routine d'avant-session. Animer une session réelle.",
     "L'offre repose enfin sur des faits et non sur ce qu'il imagine : il sait à quoi il doit être "
     "à la hauteur. Le renforcement positif est le contre-poison exact d'un 1 à 29 et d'un 3 à 30, "
     "qui minimisent tout."),

    ('8', '1h',
     "**Thème : la routine de pré-performance, l'imagerie avec l'imprévu inclus, et le cadre avant "
     "/ pendant / après.** Bilan chiffré des 3 indicateurs contre leur valeur de la séance 1, "
     "rédaction de son protocole personnel écrit (routine, phrase, valeurs, 3 signaux de bascule), "
     "plan de rechute. **EMCT : T.A.O.** sur sa prochaine session, avec une question à laquelle il "
     "ne sait pas répondre. Puis débriefing général.",
     "Aucune. La dernière séance rend l'autonomie, elle ne prescrit pas. Point de suivi optionnel "
     "à trois mois.",
     "Vérifie l'objectif, chiffres en main. J'entraîne la réaction à l'accroc et non le déroulé "
     "parfait : un 1 et un 3 s'effondrent au premier grain de sable s'ils n'ont visualisé que la "
     "réussite. Et un 8 doit partir en maîtrise, jamais en dépendance."),
]

AUTRES = [
    "**Objectif de l'accompagnement.** Je ne promets pas un résultat qui ne dépend pas de lui : "
    "comme un athlète qui prépare un championnat, je ne garantis pas le podium, je garantis que "
    "les conditions seront réunies. Le lui dire est déjà une intervention. **D'ici quatre mois, "
    "Éric aura construit et animé une offre de coaching de groupe fondée sur les besoins réellement "
    "exprimés par au moins dix prospects interrogés, aura animé au minimum deux sessions avec une "
    "note d'appréhension descendue sous 4 sur 10, et aura retrouvé deux soirées par semaine sans "
    "travail.** Trois indicateurs posés en séance 1 et suivis tout du long : l'appréhension de 0 à "
    "10 (pilier 1), le nombre de besoins vérifiés et non supposés (pilier 2), le nombre de soirées "
    "sans penser au projet (qui protège la personne).",

    "**Progression en trois piliers.** Connexion à soi, séances 1 à 4 : il est fatigué et sans "
    "voyant interne, le charger avant de le reconnecter à lui-même reviendrait à ajouter de la "
    "charge à un homme en surcharge. Connexion aux autres, séances 5 et 6 : c'est son point faible "
    "objectif et le cœur de son métier. Connexion au projet, séances 7 et 8 : l'offre ne peut se "
    "construire qu'une fois les besoins connus et les valeurs posées. Le pilier 2 ne compte que "
    "deux séances mais porte les deux actions de terrain les plus lourdes : le poids d'un pilier se "
    "mesure au travail entre les séances, pas à leur nombre.",

    "**Fréquence et durée globale.** Huit séances, une toutes les deux semaines en moyenne, "
    "l'intervalle passant à trois semaines après les séances 5 et 7, celles qui déclenchent une "
    "action de terrain lourde : le rythme suit le travail. Soit **environ quatre mois**. La séance "
    "1 dure 1h30 parce qu'elle porte l'anamnèse et l'installation de l'alliance, les sept autres 1h.",

    "**Lieu des coachings.** En visioconférence par défaut, ce qui convient à un entrepreneur en "
    "lancement et permet l'enregistrement des séances s'il le souhaite. En présentiel pour la "
    "séance 1, si possible, parce que le premier contact se joue en grande partie sur un ressenti "
    "inconscient, et pour la séance 8. Jamais à son domicile : il cache sa démarche à sa famille.",

    "**Entre les séances.** Il repart à chaque étape avec **sa roue**, qu'il retravaille et "
    "redépose, et qui sert de carnet de suivi comme de support aux échelles de 0 à 10. **La "
    "boussole de l'évolution individuelle** la complète et se repositionne aux trois charnières : "
    "après le décryptage, à l'ouverture du pilier 2, au débriefing général. J'y ajoute des **audios "
    "personnalisés enregistrés de ma voix**, protocole par protocole, pour que l'intervalle soit "
    "productif au lieu d'être creux.",

    "**Tarif : 1 600 €, payables en quatre fois de 400 €**, soit un versement par mois. Je démarre "
    "en préparation mentale mais j'ai quinze ans d'accompagnement d'entrepreneurs derrière moi, et "
    "Éric est un entrepreneur en lancement, pas un athlète : le tarif est cohérent avec ce que je "
    "vaux sur ce cas précis. Vendre n'est pas le sujet, tenir l'est. Et je peux l'annoncer sans "
    "trembler, ce qui est le critère décisif et ce que je lui enseigne en séance 7 : un prix qu'on "
    "n'assume pas s'entend à l'oral et se négocie tout seul. Il est en lancement, sa trésorerie est "
    "tendue et il cache sa démarche : le versement mensuel évite d'en faire un sujet à la maison.",

    "**Évolutivité.** Si l'hypothèse de type ne se confirme pas en séance 2, les rééquilibrages "
    "changent et les séances 5 à 7 changent de contenu, l'ossature des piliers restant la même. Si "
    "le screening révèle un épuisement clinique, j'oriente et je prolonge le pilier 1. S'il "
    "n'obtient que trois entretiens au lieu de dix, je ne rattrape pas le compte, je travaille "
    "l'obstacle, et la séance 6 arrive plus tôt. Si les entretiens révèlent un besoin très "
    "différent de son offre initiale, le pilier 3 s'alourdit d'une séance, et ce sera une bonne "
    "nouvelle. S'il refuse d'en parler à un proche, je ne force pas et je reste sur ce que le "
    "secret lui coûte.",
]


def ecrire(paragraphe, texte, taille, gras_global=False):
    """Vide un paragraphe et y ecrit le texte, ** ** delimitant le gras.

    La police n'est jamais posee : elle vient du style Normal du document, donc
    de la trame. On ne fixe que la taille.
    """
    for run in list(paragraphe.runs):
        run._element.getparent().remove(run._element)
    for i, morceau in enumerate(texte.split('**')):
        if not morceau:
            continue
        run = paragraphe.add_run(morceau)
        run.bold = gras_global or (i % 2 == 1)
        run.font.size = Pt(taille)
    pf = paragraphe.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    return paragraphe


def supprimer(paragraphe):
    """Retire le paragraphe du document.

    On ne se contente pas de vider ses runs : un paragraphe vide occupe une
    ligne, et la trame en reserve une vingtaine. Les garder ferait deborder le
    rendu au-dela des 3 pages autorisees.
    """
    element = paragraphe._element
    element.getparent().remove(element)


def remplir_cellule(cellule, texte, taille, centre=False):
    cellule.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    premiers = cellule.paragraphs
    p = premiers[0]
    for extra in premiers[1:]:
        extra._element.getparent().remove(extra._element)
    ecrire(p, texte, taille)
    if centre:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    if not os.path.exists(SOURCE):
        sys.exit('Trame introuvable : %s' % SOURCE)
    doc = Document(SOURCE)
    # On capture les paragraphes AVANT toute suppression : les index bougeraient.
    paras = list(doc.paragraphs)
    a_supprimer = []

    # 1. Les noms des candidats (paragraphes 7, 8, 9 de la trame).
    for i, nom in zip((7, 8, 9), CANDIDATS):
        ecrire(paras[i], nom, 10)

    # 2. Les axes de reflexion, dans l'espace prevu (paragraphes 14 a 21).
    ecrire(paras[14], AXES, 10)
    paras[14].paragraph_format.space_after = Pt(0)
    a_supprimer.extend(paras[15:22])

    # 3. Le tableau : la trame offre 7 lignes de donnees, il en faut 8.
    table = doc.tables[0]
    while len(table.rows) - 1 < len(SEANCES):
        table._tbl.append(copy.deepcopy(table.rows[-1]._tr))
    while len(table.rows) - 1 > len(SEANCES):
        table._tbl.remove(table.rows[-1]._tr)

    for idx, (num, duree, contenu, actions, interet) in enumerate(SEANCES, start=1):
        ligne = table.rows[idx]
        remplir_cellule(ligne.cells[0], '**%s**' % num, 11, centre=True)
        remplir_cellule(ligne.cells[1], duree, 10, centre=True)
        remplir_cellule(ligne.cells[2], contenu, 9)
        remplir_cellule(ligne.cells[3], actions, 9)
        remplir_cellule(ligne.cells[4], interet, 9)

    # 4. Autres informations, dans l'espace qui suit (paragraphes 25 et plus).
    libres = list(paras[25:])
    if len(libres) < len(AUTRES):
        sys.exit('Pas assez de paragraphes libres dans la trame.')
    for bloc, paragraphe in zip(AUTRES, libres):
        ecrire(paragraphe, bloc, 9)
    a_supprimer.extend(libres[len(AUTRES):])

    for paragraphe in a_supprimer:
        supprimer(paragraphe)

    doc.core_properties.title = 'Certification coach en préparation mentale - Thème B'
    doc.core_properties.author = 'Gisèle, Anouschka, Mac Arthur'
    doc.save(SORTIE)
    print('Ecrit :', SORTIE)


if __name__ == '__main__':
    main()
