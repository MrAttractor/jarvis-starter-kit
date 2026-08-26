# -*- coding: utf-8 -*-
"""Page HTML de document -> Word modifiable.

Usage :
  python scripts/html2docx.py source.html sortie.docx [--pied "..."]
                              [--auteur "..."] [--sans-avertissement]

Ecrit pour les pages de document de l'agence : bandeau de brouillon optionnel,
en-tete avec titre et sous-titre, puis un corps compose de h2, h3, p, ul,
blockquote, fiches en deux colonnes, tableaux, blocs de formule et pied.

Pourquoi convertir depuis le HTML et non depuis une source Markdown parallele :
les montants ne doivent exister qu'a un seul endroit. Le HTML reste la source,
le .docx est un produit derive qu'on regenere au lieu de le maintenir a la main.
"""
import re
import sys
import html as htmllib
import argparse

sys.path.insert(0, __file__.rsplit('\\', 1)[0].rsplit('/', 1)[0])
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx_commun import (nouveau_document, citation, filet, inline,
                         pied_de_page, ombrer, _ombrer_cellule, GRIS)

BLOCS = re.compile(
    r'<h2>(?P<h2>.*?)</h2>'
    r'|<h3>(?P<h3>.*?)</h3>'
    r'|<p>(?P<p>.*?)</p>'
    r'|<ul>(?P<ul>.*?)</ul>'
    r'|<blockquote>(?P<bq>.*?)</blockquote>'
    r'|<div class="fiche">(?P<fiche>.*?)</dl>\s*</div>'
    r'|<div class="tbl">(?P<tbl>.*?)</table>\s*</div>'
    r'|<div class="palier(?P<phare>[^"]*)">(?P<palier>.*?)</ul>\s*</div>'
    r'|<div class="pied">(?P<pied>.*?)</div>',
    re.S)


def texte_nu(frag):
    frag = re.sub(r'<sup>(.*?)</sup>', r'\1', frag, flags=re.S)
    frag = re.sub(r'<br\s*/?>', ' ', frag)
    frag = re.sub(r'<[^>]+>', '', frag)
    return htmllib.unescape(re.sub(r'\s+', ' ', frag)).strip()


JETONS = re.compile(r'(<strong>.*?</strong>'
                    r'|<span class="tofill">.*?</span>'
                    r'|<span class="badge">.*?</span>'
                    r'|<br\s*/?>)', re.S)


def inline_html(paragraphe, frag, base_taille=None, base_gras=False):
    """Ecrit un fragment HTML : gras conserve, champs a completer surlignes."""
    for jeton in JETONS.split(frag):
        if not jeton:
            continue
        if re.match(r'<br\s*/?>', jeton):
            paragraphe.add_run().add_break()
            continue
        gras, surligne = base_gras, False
        if jeton.startswith('<strong>'):
            gras = True
        elif jeton.startswith('<span class="tofill">') or \
                jeton.startswith('<span class="badge">'):
            surligne = True
        contenu = texte_nu(jeton)
        if not contenu:
            continue
        run = paragraphe.add_run(contenu)
        run.font.bold = gras or surligne
        if base_taille:
            run.font.size = Pt(base_taille)
        if surligne:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return paragraphe


def puces(doc, frag_ul, taille=None):
    for item in re.findall(r'<li>(.*?)</li>', frag_ul, re.S):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        inline_html(p, item, base_taille=taille)


def fiche(doc, frag):
    paires = re.findall(r'<dt>(.*?)</dt>\s*<dd>(.*?)</dd>', frag, re.S)
    if not paires:
        return
    t = doc.add_table(rows=0, cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for dt, dd in paires:
        cells = t.add_row().cells
        cells[0].width = Cm(4.2)
        cells[1].width = Cm(12)
        for cell, frag_cell, gras in ((cells[0], dt, True), (cells[1], dd, False)):
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            inline_html(p, frag_cell, base_taille=10, base_gras=gras)
        _ombrer_cellule(cells[0], 'F2EEE7')
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def table_html(doc, frag):
    entetes = re.findall(r'<th>(.*?)</th>', frag, re.S)
    lignes = [re.findall(r'<td>(.*?)</td>', tr, re.S)
              for tr in re.findall(r'<tr>(.*?)</tr>', frag, re.S)]
    lignes = [l for l in lignes if l]
    if not entetes or not lignes:
        return
    t = doc.add_table(rows=1, cols=len(entetes))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, e in enumerate(entetes):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        inline_html(p, e, base_taille=10, base_gras=True)
        _ombrer_cellule(cell, 'EAF0F3')
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    trPr = t.rows[0]._tr.get_or_add_trPr()
    rep = OxmlElement('w:tblHeader')
    rep.set(qn('w:val'), 'true')
    trPr.append(rep)
    for ligne in lignes:
        cells = t.add_row().cells
        for i, val in enumerate(ligne[:len(entetes)]):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            inline_html(p, val, base_taille=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def palier(doc, frag, phare=False):
    nom = re.search(r'<span class="nom">(.*?)</span>', frag, re.S)
    prix = re.search(r'<span class="prix">(.*?)</span>', frag, re.S)
    badge = re.search(r'<span class="badge">(.*?)</span>', frag, re.S)
    pour = re.search(r'<div class="pour">(.*?)</div>', frag, re.S)
    ul = re.search(r'<ul>(.*)', frag, re.S)

    titre = doc.add_paragraph(style='Heading 2')
    inline(titre, texte_nu(nom.group(1)) if nom else 'Formule', base_gras=True)

    ligne = doc.add_paragraph()
    ligne.paragraph_format.space_after = Pt(2)
    bits = []
    if prix:
        bits.append(texte_nu(prix.group(1)))
    if badge:
        bits.append(texte_nu(badge.group(1)))
    r = ligne.add_run(' · '.join(bits))
    r.font.bold = True
    r.font.size = Pt(11)
    ombrer(ligne, 'F7F0DF')

    if pour:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        inline_html(p, pour.group(1), base_taille=10.5)
        for run in p.runs:
            run.font.italic = True
            run.font.color.rgb = GRIS
    if ul:
        puces(doc, ul.group(1), taille=10.5)


def pied_bloc(doc, frag):
    filet(doc)
    p = doc.add_paragraph()
    inline_html(p, frag, base_taille=9.5)
    for run in p.runs:
        run.font.color.rgb = GRIS


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('source')
    ap.add_argument('sortie')
    ap.add_argument('--pied')
    ap.add_argument('--auteur', default='Thim Production')
    ap.add_argument('--sans-avertissement', action='store_true')
    a = ap.parse_args()

    src = open(a.source, encoding='utf-8').read()

    m_h1 = re.search(r'<h1>(.*?)</h1>', src, re.S)
    m_sous = re.search(r'<div class="sous">(.*?)</div>', src, re.S)
    m_avert = re.search(r'<div class="brouillon">(.*?)</div>\s*</div>', src, re.S)
    m_main = re.search(r'<main class="wrap">(.*)</main>', src, re.S)
    if not m_main:
        raise SystemExit('aucun <main class="wrap"> dans ' + a.source)

    avert = None
    if m_avert and not a.sans_avertissement:
        avert = texte_nu(m_avert.group(1))

    doc = nouveau_document(
        texte_nu(m_h1.group(1)) if m_h1 else 'Document',
        texte_nu(m_sous.group(1)) if m_sous else None,
        a.auteur, avert)

    for m in BLOCS.finditer(m_main.group(1)):
        d = m.groupdict()
        if d['h2'] is not None:
            p = doc.add_paragraph(style='Heading 1')
            inline(p, texte_nu(d['h2']), base_gras=True)
        elif d['h3'] is not None:
            p = doc.add_paragraph(style='Heading 3')
            inline(p, texte_nu(d['h3']), base_gras=True)
        elif d['p'] is not None:
            inline_html(doc.add_paragraph(), d['p'])
        elif d['ul'] is not None:
            puces(doc, d['ul'])
        elif d['bq'] is not None:
            citation(doc, texte_nu(d['bq']))
        elif d['fiche'] is not None:
            fiche(doc, d['fiche'])
        elif d['tbl'] is not None:
            table_html(doc, d['tbl'])
        elif d['palier'] is not None:
            palier(doc, d['palier'], d['phare'])
        elif d['pied'] is not None:
            pied_bloc(doc, d['pied'])

    if a.pied:
        pied_de_page(doc, a.pied)
    doc.save(a.sortie)
    print('ecrit', a.sortie)


if __name__ == '__main__':
    main()
