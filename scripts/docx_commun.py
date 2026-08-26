# -*- coding: utf-8 -*-
"""Briques communes pour produire des documents Word modifiables par un client.

Pourquoi ce module existe : un client qui doit travailler sur un livrable ne peut
rien faire d'un fichier Markdown ni d'une page HTML. Il lui faut un .docx qu'il
ouvre, modifie et renvoie. Ces briques sont partagees par md2docx.py et
html2docx.py pour que les deux sorties se ressemblent.

Choix assumes :
  - police Calibri 11, celle par defaut de Word, pour que le client ne se batte
    pas avec une police absente de son poste ;
  - langue fr-FR posee sur les styles, sinon le correcteur souligne tout ;
  - les champs a completer sont surlignes en jaune, impossible a rater a l'ecran
    comme a l'impression ;
  - l'auteur des proprietes du fichier est celui qui envoie le document, pas
    l'agence qui le redige : ces pieces partent chez des tiers.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX
import re

ARDOISE = RGBColor(0x2F, 0x48, 0x58)
ENCRE = RGBColor(0x1A, 0x1A, 0x18)
GRIS = RGBColor(0x84, 0x83, 0x7C)
ROUGE = RGBColor(0x8C, 0x3A, 0x32)
FOND_BLOC = 'F7F6F4'


def _langue(style, code='fr-FR'):
    rpr = style.element.get_or_add_rPr()
    for balise in rpr.findall(qn('w:lang')):
        rpr.remove(balise)
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), code)
    rpr.append(lang)


def nouveau_document(titre, sous_titre=None, auteur='Thim Production',
                     avertissement=None):
    doc = Document()

    doc.core_properties.title = titre
    doc.core_properties.author = auteur
    doc.core_properties.last_modified_by = auteur
    doc.core_properties.comments = (
        'Document de travail modifiable. Festival des Grillades de Paris 2026.')

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.font.color.rgb = ENCRE
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15
    _langue(normal)

    for nom, taille, couleur, avant in (
            ('Title', 22, ARDOISE, 0), ('Heading 1', 15, ARDOISE, 16),
            ('Heading 2', 12.5, ARDOISE, 12), ('Heading 3', 11.5, ENCRE, 10)):
        st = doc.styles[nom]
        st.font.name = 'Calibri'
        st.font.size = Pt(taille)
        st.font.bold = True
        st.font.color.rgb = couleur
        st.font.italic = False
        st.paragraph_format.space_before = Pt(avant)
        st.paragraph_format.space_after = Pt(4)
        _langue(st)

    p = doc.add_paragraph(titre, style='Title')
    p.paragraph_format.space_after = Pt(2)
    if sous_titre:
        s = doc.add_paragraph()
        r = s.add_run(sous_titre)
        r.font.size = Pt(11.5)
        r.font.color.rgb = GRIS
        s.paragraph_format.space_after = Pt(14)

    if avertissement:
        bloc = doc.add_paragraph()
        run = bloc.add_run(avertissement)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = ROUGE
        ombrer(bloc, 'FBEEEC')
        bloc.paragraph_format.space_before = Pt(4)
        bloc.paragraph_format.space_after = Pt(14)
        encadrer(bloc, '8C3A32')

    return doc


def ombrer(paragraphe, couleur_fond):
    pPr = paragraphe._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), couleur_fond)
    pPr.append(shd)


def encadrer(paragraphe, couleur='D2CFC8'):
    pPr = paragraphe._p.get_or_add_pPr()
    bd = OxmlElement('w:pBdr')
    for cote in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement('w:' + cote)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '6')
        el.set(qn('w:color'), couleur)
        bd.append(el)
    pPr.append(bd)


# Un champ a completer se reconnait a ses crochets. Il est surligne, pas ecrit
# en petit dans une note de bas de page : c'est la seule facon qu'il ne parte
# pas tel quel chez un sponsor.
CHAMP = re.compile(r'\[[^\[\]]{2,120}\]')


def inline(paragraphe, texte, base_gras=False, base_taille=None):
    """Ecrit du texte en gerant **gras**, `code`, *italique* et [champs]."""
    morceaux = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*\n]+\*|\[[^\[\]]{2,120}\])',
                        texte)
    for m in morceaux:
        if not m:
            continue
        gras, italique, mono, champ = base_gras, False, False, False
        if m.startswith('**') and m.endswith('**') and len(m) > 4:
            m, gras = m[2:-2], True
        elif m.startswith('`') and m.endswith('`') and len(m) > 2:
            m, mono = m[1:-1], True
        elif m.startswith('*') and m.endswith('*') and len(m) > 2:
            m, italique = m[1:-1], True
        elif CHAMP.fullmatch(m):
            champ = True
        run = paragraphe.add_run(m)
        run.font.bold = gras
        run.font.italic = italique
        if base_taille:
            run.font.size = Pt(base_taille)
        if mono:
            run.font.name = 'Consolas'
            run.font.size = Pt(base_taille or 10)
        if champ:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            run.font.bold = True
    return paragraphe


def bloc_a_copier(doc, lignes):
    """Un texte pret a etre copie : lettre type, message, script d'appel.

    Volontairement dans la police du document et non en machine a ecrire : le
    client va le reutiliser tel quel. Une ligne de la source donne un paragraphe
    Word, les lignes vides sont ignorees et l'espacement est porte par le style.
    C'est ce qui permet au client de reecrire une phrase sans que le texte reste
    coupe la ou la source l'avait coupe : dans la source, un paragraphe de lettre
    tient donc sur une seule ligne, aussi longue qu'il faut.
    """
    premier = None
    for ligne in lignes:
        if not ligne.strip():
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.right_indent = Cm(0.3)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(6 if premier is None else 0)
        p.paragraph_format.line_spacing = 1.15
        inline(p, ligne.strip(), base_taille=10.5)
        ombrer(p, FOND_BLOC)
        premier = premier or p
    return premier


def tableau(doc, entetes, lignes):
    t = doc.add_table(rows=1, cols=len(entetes))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True

    for i, e in enumerate(entetes):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        inline(p, e, base_gras=True, base_taille=10)
        _ombrer_cellule(cell, 'EAF0F3')

    tr = t.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    entete = OxmlElement('w:tblHeader')
    entete.set(qn('w:val'), 'true')
    trPr.append(entete)

    for ligne in lignes:
        cells = t.add_row().cells
        for i, valeur in enumerate(ligne[:len(entetes)]):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            inline(p, valeur, base_taille=10)
    return t


def _ombrer_cellule(cell, couleur):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), couleur)
    tcPr.append(shd)


def citation(doc, texte):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    inline(p, texte, base_taille=10.5)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x3D, 0x3D, 0x39)
    ombrer(p, 'FFFDF5')
    encadrer(p, 'C9A54E')
    return p


def filet(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    bd = OxmlElement('w:pBdr')
    el = OxmlElement('w:bottom')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), '6')
    el.set(qn('w:space'), '1')
    el.set(qn('w:color'), 'D2CFC8')
    bd.append(el)
    pPr.append(bd)
    return p


def pied_de_page(doc, texte):
    section = doc.sections[0]
    p = section.footer.paragraphs[0]
    p.text = ''
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texte)
    run.font.size = Pt(8.5)
    run.font.color.rgb = GRIS
